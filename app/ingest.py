import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path

from docx import Document as DocxDocument
from langchain_community.document_loaders import BSHTMLLoader, PyPDFLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from app.settings import CHUNK_OVERLAP, CHUNK_SIZE, MANIFEST_PATH


def file_sha256(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def _load_docx(path: Path) -> list:
    doc = DocxDocument(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    text = "\n\n".join(parts)
    return [Document(page_content=text, metadata={"source": str(path), "page": 0})]


def load_file(path: Path) -> list:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return PyPDFLoader(str(path)).load()
    if suffix in {".html", ".htm"}:
        return BSHTMLLoader(str(path)).load()
    if suffix == ".docx":
        return _load_docx(path)
    raise ValueError(f"Unsupported file type: {suffix}")


def chunk_documents(docs: list, file_name: str, sha256: str) -> list[dict]:
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE, chunk_overlap=CHUNK_OVERLAP
    )
    chunks = splitter.split_documents(docs)
    now = datetime.now(timezone.utc).isoformat()
    result = []
    for i, chunk in enumerate(chunks):
        page = chunk.metadata.get("page", 0)
        cid = f"{sha256[:12]}_{page}_{i}"
        result.append({
            "chunk_id": cid,
            "content": chunk.page_content,
            "metadata": {
                "source": file_name,
                "file_name": file_name,
                "file_type": Path(file_name).suffix.lstrip("."),
                "page": page,
                "chunk_id": cid,
                "sha256": sha256,
                "uploaded_by_user": True,
                "uploaded_at": now,
            },
        })
    return result


def upsert_chunks(chunks: list[dict]) -> None:
    # Import here to avoid circular import at module load time
    from app import rag as _rag

    ids        = [c["chunk_id"] for c in chunks]
    docs       = [c["content"]  for c in chunks]
    metas      = [c["metadata"] for c in chunks]
    embeddings = _rag._model.encode(docs).tolist()
    _rag._collection.upsert(ids=ids, embeddings=embeddings, documents=docs, metadatas=metas)


def _load_manifest() -> dict:
    p = Path(MANIFEST_PATH)
    return json.loads(p.read_text()) if p.exists() else {}


def _save_manifest(manifest: dict) -> None:
    Path(MANIFEST_PATH).write_text(json.dumps(manifest, indent=2))


def is_already_indexed(sha256: str) -> bool:
    return any(
        v.get("content_hash") == sha256
        for v in _load_manifest().values()
    )


class IngestReport:
    def __init__(self, sha256: str, chunk_count: int, collection_size: int):
        self.sha256 = sha256
        self.chunk_count = chunk_count
        self.collection_size = collection_size


def ingest_file(path: Path) -> IngestReport:
    from app import rag as _rag

    sha256 = file_sha256(path)
    docs   = load_file(path)
    chunks = chunk_documents(docs, path.name, sha256)
    upsert_chunks(chunks)

    manifest = _load_manifest()
    manifest[f"upload/{path.name}"] = {
        "content_hash":   sha256,
        "chunk_count":    len(chunks),
        "last_processed": datetime.now(timezone.utc).isoformat(),
    }
    _save_manifest(manifest)

    return IngestReport(
        sha256=sha256,
        chunk_count=len(chunks),
        collection_size=_rag._collection.count(),
    )
