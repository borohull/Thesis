"""
Build or incrementally update the ChromaDB knowledge base.

Run from the project root:
    python scripts/build_index.py

Combines the logic of notebooks/01_data_ingestion.ipynb and
notebooks/02_embedding.ipynb into a single executable script.
"""
import warnings
warnings.filterwarnings("ignore", category=UserWarning)
import hashlib
import json
import sys
from datetime import datetime
from pathlib import Path
from typing import List

import numpy as np
import pymupdf
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langchain_community.document_loaders import TextLoader
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
import chromadb

# All paths and config from settings
sys.path.insert(0, str(Path(__file__).parent.parent))
from app.settings import (
    RAW_DIR,
    CHUNKS_PATH,
    MANIFEST_PATH,
    PENDING_CHANGES_PATH,
    CHROMA_PATH,
    COLLECTION_NAME,
    EMBEDDING_MODEL,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
)

# Constants

SUPPORTED_EXTS = {".pdf", ".txt", ".md", ".html", ".htm", ".docx"}
MIN_CHUNK_SIZE = 100

# Text loading functions (verbatim from notebooks/01_data_ingestion.ipynb)

def clean_text(text: str) -> str:
    import re
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def load_pdf_pymupdf(file_path) -> List[Document]:
    file_path = Path(file_path)
    pdf = pymupdf.open(str(file_path))

    linked_from = file_path.parent.name if "linked_pdfs" in file_path.parts else None

    docs = []
    for page_num, page in enumerate(pdf):
        text = page.get_text("text")
        docs.append(
            Document(
                page_content=clean_text(text),
                metadata={
                    "source": str(file_path),
                    "file_name": file_path.name,
                    "file_type": "pdf",
                    "page": page_num,
                    **({"linked_from": linked_from} if linked_from else {})
                }
            )
        )

    pdf.close()
    return docs


def load_txt(file_path, encoding: str = "utf-8") -> List[Document]:
    file_path = Path(file_path)
    loader = TextLoader(str(file_path), encoding=encoding)
    docs = loader.load()

    normalized_docs = []
    for doc in docs:
        normalized_docs.append(
            Document(
                page_content=clean_text(doc.page_content),
                metadata={
                    "source": str(file_path),
                    "file_name": file_path.name,
                    "file_type": "txt",
                    **doc.metadata
                }
            )
        )

    return normalized_docs


def load_html(file_path) -> List[Document]:
    import re as _re
    file_path = Path(file_path)

    with open(file_path, "r", encoding="utf-8") as f:
        html = f.read()

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "nav", "footer", "header", "aside"]):
        tag.decompose()

    # Wider content-div fallback
    root = (
        soup.find("main")
        or soup.find("article")
        or soup.find("div", id=_re.compile(r"(main|content|page)", _re.I))
        or soup.find("div", class_=_re.compile(r"(main|content|page)", _re.I))
        or soup.body
        or soup
    )

    # Page title for context prefix
    title_str = soup.title.string.strip() if soup.title and soup.title.string else ""
    
    meta_desc = ""
    meta_tag = soup.find("meta", attrs={"name": "description"})
    if meta_tag and meta_tag.get("content", "").strip():
        meta_desc = meta_tag["content"].strip()

    current_heading: str = ""
    lines: list[str] = []

    for tag in root.find_all(["h1", "h2", "h3", "h4", "h5", "h6",
                               "p", "li", "td", "dd"]):
        text = tag.get_text(separator=" ", strip=True)
        if not text:
            continue
        if tag.name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            current_heading = text
        else:
            lines.append(f"[{current_heading}] {text}" if current_heading else text)

    # Filter boilerplate
    lines = [l for l in lines if len(l.split()) >= 4]

    body = "\n".join(lines)
    parts = [p for p in [title_str, meta_desc, body] if p]
    text = "\n\n".join(parts)

    linked_from = file_path.parent.name if "linked_pdfs" in file_path.parts else None

    return [
        Document(
            page_content=clean_text(text),
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "html",
                "title": title_str or None,
                **({"linked_from": linked_from} if linked_from else {})
            }
        )
    ]


def load_docx(file_path) -> List[Document]:
    file_path = Path(file_path)
    doc = DocxDocument(str(file_path))

    if all(not p.text.strip() for p in doc.paragraphs):
        print(f"  Warning: {file_path.name} has no paragraph content, skipping")
        return []

    current_heading: str = ""
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style.name.startswith("Heading"):
            current_heading = text
        else:
            lines.append(f"[{current_heading}] {text}" if current_heading else text)

    for table in doc.tables:
        for row in table.rows:
            cells = " | ".join(cell.text.strip() for cell in row.cells)
            if cells.replace("|", "").strip():
                lines.append(f"[{current_heading}] {cells}" if current_heading else cells)

    linked_from = file_path.parent.name if "linked_docx" in file_path.parts else None

    return [
        Document(
            page_content=clean_text("\n".join(lines)),
            metadata={
                "source": str(file_path),
                "file_name": file_path.name,
                "file_type": "docx",
                **({"linked_from": linked_from} if linked_from else {})
            }
        )
    ]


def load_file(file_path) -> List[Document]:
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return load_pdf_pymupdf(file_path)
    elif suffix in {".txt", ".md"}:
        return load_txt(file_path)
    elif suffix in {".html", ".htm"}:
        return load_html(file_path)
    elif suffix == ".docx":
        return load_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


# Manifest / chunk helpers (verbatim from notebooks/01_data_ingestion.ipynb)

def compute_file_hash(file_path) -> str:
    h = hashlib.sha256()
    with open(file_path, "rb") as f:
        for block in iter(lambda: f.read(65536), b""):
            h.update(block)
    return h.hexdigest()


def get_relative_path(file_path, raw_dir) -> str:
    return Path(file_path).resolve().relative_to(Path(raw_dir).resolve()).as_posix()


def load_manifest(path) -> dict:
    path = Path(path)
    if not path.exists():
        return {}
    with open(path, encoding="utf-8") as f:
        return json.load(f)


def save_manifest(manifest: dict, path) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)


def save_chunks(chunks_list: list, path) -> None:
    Path(path).parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(chunks_list, f, ensure_ascii=False, indent=2)


def load_chunks(path) -> list:
    path = Path(path)
    if not path.exists():
        return []
    with open(path, encoding="utf-8") as f:
        return json.load(f)


# EmbeddingPipeline (verbatim from notebooks/02_embedding.ipynb)

class EmbeddingPipeline:
    def __init__(self, model_name: str = EMBEDDING_MODEL):
        self.model = SentenceTransformer(model_name)
        print(f"Loaded model: {model_name}")

    def encode(self, texts: list) -> np.ndarray:
        return self.model.encode(texts, show_progress_bar=True)


# Main pipeline

if __name__ == "__main__":
    print("=== build_index.py: ELTE IK knowledge base builder ===\n")

    # Phase 1: Ingestion (notebook 01 logic)

    print("--- Phase 1: Ingestion ---")

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""],
    )

    def chunk_file(file_path: Path, content_hash: str, source_relative: str) -> list:
        docs = load_file(file_path)
        if not docs:
            return []
        split_docs = text_splitter.split_documents(docs)
        split_docs = [d for d in split_docs if len(d.page_content) >= MIN_CHUNK_SIZE]
        chunks_out = []
        for n, doc in enumerate(split_docs):
            meta = dict(doc.metadata)
            meta["source_relative"] = source_relative
            meta["content_hash"]    = content_hash
            meta["chunk_id"]        = f"{source_relative}::chunk_{n}"
            chunks_out.append({"content": doc.page_content, "metadata": meta})
        return chunks_out

    # 1. Scan disk
    raw_dir = Path(RAW_DIR)
    if not raw_dir.exists():
        print(f"ERROR: RAW_DIR does not exist: {RAW_DIR}")
        sys.exit(1)

    current_files = {}
    for file_path in raw_dir.rglob("*"):
        if not file_path.is_file():
            continue
        if file_path.suffix.lower() not in SUPPORTED_EXTS:
            continue
        rel = get_relative_path(file_path, raw_dir)
        current_files[rel] = (file_path, compute_file_hash(file_path))

    # 2. Load manifest and compute diff
    manifest = load_manifest(MANIFEST_PATH)
    manifest_paths = set(manifest.keys())
    disk_paths     = set(current_files.keys())

    new_files     = sorted(disk_paths - manifest_paths)
    deleted_files = sorted(manifest_paths - disk_paths)
    updated_files = sorted(
        p for p in (disk_paths & manifest_paths)
        if manifest[p]["content_hash"] != current_files[p][1]
    )
    unchanged_count = len((disk_paths & manifest_paths) - set(updated_files))

    print(f"Scan: {len(disk_paths)} files on disk, {len(manifest_paths)} in manifest")
    print(f"  new:       {len(new_files)}")
    print(f"  updated:   {len(updated_files)}")
    print(f"  deleted:   {len(deleted_files)}")
    print(f"  unchanged: {unchanged_count}")

    # 3. Load existing chunks.json and remove entries for deleted/updated files
    existing_chunks = load_chunks(CHUNKS_PATH)
    to_remove = set(deleted_files) | set(updated_files)
    kept_chunks = [
        c for c in existing_chunks
        if c["metadata"].get("source_relative")
        and c["metadata"].get("source_relative") not in to_remove
    ]
    removed_count = len(existing_chunks) - len(kept_chunks)
    if removed_count:
        print(f"Removed {removed_count} chunks from chunks.json (deleted/updated/old-schema)")

    # 4. Chunk new + updated files
    to_process = new_files + updated_files
    new_chunks_total = []
    for rel in to_process:
        file_path, content_hash = current_files[rel]
        try:
            file_chunks = chunk_file(file_path, content_hash, rel)
            new_chunks_total.extend(file_chunks)
            manifest[rel] = {
                "content_hash":   content_hash,
                "chunk_count":    len(file_chunks),
                "last_processed": datetime.now().isoformat(timespec="seconds"),
            }
            print(f"  Chunked {rel}: {len(file_chunks)} chunks")
        except Exception as e:
            print(f"  Failed {rel}: {e}")

    # 5. Drop deleted files from manifest
    for rel in deleted_files:
        if rel in manifest:
            del manifest[rel]

    # 6. Persist chunks.json, manifest, pending_changes
    all_chunks = kept_chunks + new_chunks_total
    save_chunks(all_chunks, CHUNKS_PATH)
    save_manifest(manifest, MANIFEST_PATH)

    pending_changes = {
        "new":     new_files,
        "updated": updated_files,
        "deleted": deleted_files,
    }
    pending_path = Path(PENDING_CHANGES_PATH)
    pending_path.parent.mkdir(parents=True, exist_ok=True)
    with open(pending_path, "w", encoding="utf-8") as f:
        json.dump(pending_changes, f, ensure_ascii=False, indent=2)

    print(f"\nTotal chunks in chunks.json: {len(all_chunks)}")

    # Phase 2: Embedding (notebook 02 logic)

    print("\n--- Phase 2: Embedding ---")

    with open(CHUNKS_PATH, encoding="utf-8") as f:
        chunks = json.load(f)
    print(f"Loaded {len(chunks)} chunks")

    if pending_path.exists():
        with open(pending_path, encoding="utf-8") as f:
            pending = json.load(f)
        print(f"Pending changes: {len(pending['new'])} new, "
              f"{len(pending['updated'])} updated, "
              f"{len(pending['deleted'])} deleted")
    else:
        pending = {
            "new":     sorted({c["metadata"].get("source_relative") for c in chunks
                               if c["metadata"].get("source_relative")}),
            "updated": [],
            "deleted": [],
        }
        print(f"No pending_changes.json found -> full re-embed of {len(pending['new'])} files")

    to_embed_paths = set(pending["new"]) | set(pending["updated"])
    chunks_to_embed = [c for c in chunks if c["metadata"].get("source_relative") in to_embed_paths]
    print(f"Will embed {len(chunks_to_embed)} chunks from {len(to_embed_paths)} files")

    if chunks_to_embed:
        try:
            pipeline = EmbeddingPipeline()
            texts = [c["content"] for c in chunks_to_embed]
            embeddings = pipeline.encode(texts)
            print(f"Embeddings shape: {embeddings.shape}")
        except Exception as e:
            print(f"ERROR: Embedding failed: {e}")
            sys.exit(1)
    else:
        pipeline = None
        texts = []
        embeddings = None
        print("Nothing to embed.")

    # Open ChromaDB collection
    try:
        client = chromadb.PersistentClient(path=CHROMA_PATH)
        collection = client.get_or_create_collection(
            name=COLLECTION_NAME,
            metadata={"hnsw:space": "cosine"},
        )
    except Exception as e:
        print(f"ERROR: Could not open ChromaDB at {CHROMA_PATH}: {e}")
        sys.exit(1)

    # Delete old chunks for updated/deleted files
    ids_to_delete = []
    files_to_clear = list(set(pending["updated"]) | set(pending["deleted"]))
    for src_rel in files_to_clear:
        existing = collection.get(where={"source_relative": src_rel}, include=[])
        if existing["ids"]:
            ids_to_delete.extend(existing["ids"])

    if ids_to_delete:
        BATCH_SIZE = 5000
        for i in range(0, len(ids_to_delete), BATCH_SIZE):
            collection.delete(ids=ids_to_delete[i:i + BATCH_SIZE])
        print(f"Deleted {len(ids_to_delete)} old chunks from ChromaDB")
    else:
        print("No old chunks to delete")

    # Upsert new + updated chunks
    if chunks_to_embed:
        ids   = [c["metadata"]["chunk_id"] for c in chunks_to_embed]
        metas = [c["metadata"] for c in chunks_to_embed]
        embs  = embeddings.tolist()

        BATCH_SIZE = 5000
        for i in range(0, len(ids), BATCH_SIZE):
            collection.upsert(
                ids=ids[i:i + BATCH_SIZE],
                embeddings=embs[i:i + BATCH_SIZE],
                documents=texts[i:i + BATCH_SIZE],
                metadatas=metas[i:i + BATCH_SIZE],
            )
        print(f"Upserted {len(ids)} chunks into '{COLLECTION_NAME}'")

    print(f"Collection now has {collection.count()} documents")

    # Clean up pending_changes
    if pending_path.exists():
        pending_path.unlink()
        print(f"Removed pending_changes.json")

    print("\n=== Done. Knowledge base is up to date. ===")
